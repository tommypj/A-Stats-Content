"""
Document processor for extracting and chunking text from various file types.

Supports PDF, TXT, Markdown, DOCX, and HTML files.
"""

import logging
import re
from dataclasses import dataclass
from enum import StrEnum
from pathlib import Path
from typing import Any, BinaryIO

logger = logging.getLogger(__name__)


class DocumentProcessingError(Exception):
    """Base exception for document processing errors."""

    pass


class UnsupportedDocumentError(DocumentProcessingError):
    """Raised when document type is not supported."""

    pass


class DocumentType(StrEnum):
    """Supported document types."""

    PDF = "pdf"
    TXT = "txt"
    MARKDOWN = "md"
    DOCX = "docx"
    HTML = "html"


@dataclass
class ProcessedChunk:
    """Represents a chunk of text from a document."""

    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: dict[str, Any]


@dataclass
class ProcessedDocument:
    """Result of document processing."""

    source_id: str
    title: str
    document_type: DocumentType
    chunks: list[ProcessedChunk]
    total_chars: int
    total_chunks: int


class DocumentProcessor:
    """
    Process documents into chunks for RAG.

    Supports PDF, TXT, Markdown, DOCX, and HTML files.
    """

    DEFAULT_CHUNK_SIZE = 1000  # characters
    DEFAULT_CHUNK_OVERLAP = 200

    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        """
        Initialize document processor.

        Args:
            chunk_size: Target size for each chunk (in characters)
            chunk_overlap: Overlap between chunks to maintain context
        """
        self.chunk_size = chunk_size or self.DEFAULT_CHUNK_SIZE
        self.chunk_overlap = chunk_overlap or self.DEFAULT_CHUNK_OVERLAP

    async def process_file(
        self,
        file: BinaryIO,
        filename: str,
        source_id: str,
        metadata: dict | None = None,
    ) -> ProcessedDocument:
        """
        Process a file into chunks.

        Args:
            file: File object (BinaryIO)
            filename: Original filename
            source_id: Unique source identifier
            metadata: Optional additional metadata

        Returns:
            ProcessedDocument with chunks

        Raises:
            ValueError: If file type is not supported
        """
        doc_type = self.detect_type(filename)
        title = metadata.get("title", filename) if metadata else filename

        # Extract text based on file type
        if doc_type == DocumentType.PDF:
            text = await self.extract_text_pdf(file)
        elif doc_type == DocumentType.DOCX:
            text = await self.extract_text_docx(file)
        elif doc_type == DocumentType.HTML:
            text = await self.extract_text_html(file)
        elif doc_type in (DocumentType.TXT, DocumentType.MARKDOWN):
            text = file.read().decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file type: {doc_type}")

        # Chunk the text
        chunks = self.chunk_text(text)

        # Add metadata to each chunk
        base_metadata = metadata or {}
        base_metadata.update({"source_id": source_id, "title": title, "type": doc_type.value})

        processed_chunks = []
        for chunk in chunks:
            chunk_metadata = {**base_metadata, "chunk_index": chunk.chunk_index}
            processed_chunks.append(
                ProcessedChunk(
                    content=chunk.content,
                    chunk_index=chunk.chunk_index,
                    start_char=chunk.start_char,
                    end_char=chunk.end_char,
                    metadata=chunk_metadata,
                )
            )

        logger.info(
            f"Processed {filename} ({doc_type.value}) into {len(processed_chunks)} chunks "
            f"({len(text)} chars)"
        )

        return ProcessedDocument(
            source_id=source_id,
            title=title,
            document_type=doc_type,
            chunks=processed_chunks,
            total_chars=len(text),
            total_chunks=len(processed_chunks),
        )

    def detect_type(self, filename: str) -> DocumentType:
        """
        Detect document type from filename.

        Args:
            filename: Filename to analyze

        Returns:
            DocumentType enum

        Raises:
            ValueError: If file type is not supported
        """
        ext = Path(filename).suffix.lower().lstrip(".")

        type_mapping = {
            "pdf": DocumentType.PDF,
            "txt": DocumentType.TXT,
            "md": DocumentType.MARKDOWN,
            "markdown": DocumentType.MARKDOWN,
            "docx": DocumentType.DOCX,
            "html": DocumentType.HTML,
            "htm": DocumentType.HTML,
        }

        if ext not in type_mapping:
            raise ValueError(f"Unsupported file extension: .{ext}")

        return type_mapping[ext]

    async def extract_text_pdf(self, file: BinaryIO) -> str:
        """
        Extract text from PDF using pypdf.

        Args:
            file: PDF file object

        Returns:
            Extracted text
        """
        try:
            from pypdf import PdfReader

            reader = PdfReader(file)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            return "\n\n".join(text_parts)

        except ImportError:
            logger.error("pypdf not installed. Install with: pip install pypdf")
            raise RuntimeError("pypdf required for PDF processing")
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {e}")
            raise

    async def extract_text_docx(self, file: BinaryIO) -> str:
        """
        Extract text from DOCX using python-docx.

        Args:
            file: DOCX file object

        Returns:
            Extracted text
        """
        try:
            from docx import Document

            doc = Document(file)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)

        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            raise RuntimeError("python-docx required for DOCX processing")
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {e}")
            raise

    async def extract_text_html(self, file: BinaryIO) -> str:
        """
        Extract text from HTML using BeautifulSoup.

        Args:
            file: HTML file object

        Returns:
            Extracted text
        """
        try:
            from bs4 import BeautifulSoup

            html_content = file.read().decode("utf-8", errors="ignore")
            soup = BeautifulSoup(html_content, "html.parser")

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text
            text = soup.get_text(separator="\n")

            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = "\n".join(chunk for chunk in chunks if chunk)

            return text

        except ImportError:
            logger.error("beautifulsoup4 not installed. Install with: pip install beautifulsoup4")
            raise RuntimeError("beautifulsoup4 required for HTML processing")
        except Exception as e:
            logger.error(f"Failed to extract HTML text: {e}")
            raise

    def chunk_text(self, text: str) -> list[ProcessedChunk]:
        """
        Split text into overlapping chunks.

        Uses sentence boundaries when possible for better semantic coherence.

        Args:
            text: Text to chunk

        Returns:
            List of ProcessedChunk objects
        """
        chunks = []
        chunk_index = 0

        # Clean up excessive whitespace
        text = re.sub(r"\s+", " ", text).strip()

        if not text:
            return chunks

        start = 0
        while start < len(text):
            # Calculate end position
            end = min(start + self.chunk_size, len(text))

            # If this isn't the last chunk, try to break at a sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                search_start = max(start, end - 100)
                search_end = min(len(text), end + 100)
                search_text = text[search_start:search_end]

                # Find sentence boundaries (. ! ? followed by space or newline)
                sentence_endings = [m.end() for m in re.finditer(r"[.!?]\s+", search_text)]

                if sentence_endings:
                    # Use the sentence ending closest to our target
                    best_ending = min(sentence_endings, key=lambda x: abs(x - (end - search_start)))
                    end = search_start + best_ending

            # Extract chunk
            chunk_text = text[start:end].strip()

            if chunk_text:
                chunks.append(
                    ProcessedChunk(
                        content=chunk_text,
                        chunk_index=chunk_index,
                        start_char=start,
                        end_char=end,
                        metadata={},
                    )
                )
                chunk_index += 1

            # Move to next chunk with overlap
            start = end - self.chunk_overlap

            # Prevent infinite loop
            if start >= len(text):
                break

        return chunks


# Singleton instance
document_processor = DocumentProcessor()
